import os
from pathlib import Path
from typing import List
from langchain_core.documents import Document
import PyPDF2
import docx

def load_file(file_path: str) -> List[Document]:
    suffix = Path(file_path).suffix.lower()
    documents = []

    if suffix == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
            documents.append(Document(page_content=text, metadata={"source": file_path}))

    elif suffix == ".pdf":
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        documents.append(Document(page_content=text, metadata={"source": file_path}))

    elif suffix == ".docx":
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        documents.append(Document(page_content=text, metadata={"source": file_path}))

    elif suffix == ".md":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
            documents.append(Document(page_content=text, metadata={"source": file_path}))

    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    return documents
